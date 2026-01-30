import spacy
import fitz  # PyMuPDF
from docx import Document
import re

class RedactorEngine:
    def __init__(self, model_name="en_core_web_sm"):
        try:
            self.nlp = spacy.load(model_name)
        except OSError:
            # Fallback or error if model isn't present.
            # We assume user has it or we might need to download it (but we are offline restricted).
            # For now, let's try to load 'en_core_web_sm' or 'en_core_web_trf' if available.
            try:
                self.nlp = spacy.load("en_core_web_trf")
            except OSError:
                 raise RuntimeError("spaCy model not found. Please install 'en_core_web_sm'.")

    def identify_pii(self, text):
        doc = self.nlp(text)
        pii_spans = []
        for ent in doc.ents:
            if ent.label_ in ["PERSON", "ORG", "GPE", "DATE", "MONEY", "CARDINAL"]: 
                # Broad categories for 'Legal' context, usually includes Names, Dates, Money.
                # Adjusting based on user request "Names, SSNs, Addresses"
                # SSNs are often shaped, not just NER. 
                if ent.label_ in ["PERSON", "GPE"]:
                    pii_spans.append(ent)
        
        # Regex for SSN
        ssn_pattern = r"\b\d{3}-\d{2}-\d{4}\b"
        for match in re.finditer(ssn_pattern, text):
            # Create a span-like object or just return ranges? 
            # Integrating with spaCy spans is cleaner if we use the same structure.
            # But for simplicity, we'll return a list of text to redact or char offsets.
            # Our PDF/Docx redactors usually need text strings or regexes to find coordinates.
            pass
            
        return pii_spans

    def redact_pdf(self, input_path, output_path):
        doc = fitz.open(input_path)
        for page in doc:
            # 1. Get text to find PII
            text = page.get_text()
            spacy_doc = self.nlp(text)
            
            # 2. Identify PII entities
            redaction_list = []
            
            # NER based
            for ent in spacy_doc.ents:
                if ent.label_ in ["PERSON", "GPE", "ORG"]:
                    redaction_list.append(ent.text)
            
            # Regex based (SSN)
            ssn_matches = re.findall(r"\b\d{3}-\d{2}-\d{4}\b", text)
            redaction_list.extend(ssn_matches)
            
            # 3. Search and Redact
            # Use set to avoid duplicates
            for term in set(redaction_list):
                # Search for the term on the page
                quads = page.search_for(term)
                
                for quad in quads:
                    # Add redaction annotation (black box)
                    page.add_redact_annot(quad, fill=(0, 0, 0))
            
            # Apply redactions
            page.apply_redactions()
            
        doc.save(output_path)
        doc.close()

    def redact_docx(self, input_path, output_path):
        doc = Document(input_path)
        
        def redact_text(text):
            spacy_doc = self.nlp(text)
            modified_text = text
            
            # We replace from end to start to not mess up indices if we were using indices.
            # But simple replace works if tokens are unique enough. 
            # A safer way is to identifying ranges and building a new string.
            
            # Let's collect all intervals to redact
            intervals = []
            for ent in spacy_doc.ents:
                if ent.label_ in ["PERSON", "GPE", "ORG"]:
                    intervals.append((ent.start_char, ent.end_char))
            
            # Regex SSNs
            for match in re.finditer(r"\b\d{3}-\d{2}-\d{4}\b", text):
                intervals.append(match.span())
            
            # Sort and merge intervals
            intervals.sort(key=lambda x: x[0])
            merged = []
            if intervals:
                curr_start, curr_end = intervals[0]
                for next_start, next_end in intervals[1:]:
                    if next_start < curr_end:
                        curr_end = max(curr_end, next_end)
                    else:
                        merged.append((curr_start, curr_end))
                        curr_start, curr_end = next_start, next_end
                merged.append((curr_start, curr_end))
            
            # Reconstruct string
            last_idx = 0
            new_text = ""
            for start, end in merged:
                new_text += text[last_idx:start]
                new_text += "[REDACTED]"
                last_idx = end
            new_text += text[last_idx:]
            
            return new_text

        for para in doc.paragraphs:
            if para.text.strip():
                para.text = redact_text(para.text)
                
        # Handle tables? For MVP, paragraphs are usually enough.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            para.text = redact_text(para.text)

        doc.save(output_path)
