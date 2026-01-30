from docx import Document

doc = Document()
doc.add_heading('Confidential Document', 0)
doc.add_paragraph('This is a sensitive legal document.')
doc.add_paragraph('My name is John Doe and I live in New York.')
doc.add_paragraph('My social security number is 123-45-6789.')
doc.add_paragraph('Please redact this information.')

doc.save('sample_with_pii.docx')
print("Created sample_with_pii.docx")
